# -*- coding: utf-8 -*-
import chardet
from collections import namedtuple, OrderedDict
import bibtexparser
import docx
from docx.shared import Inches
from docx.enum.text import WD_BREAK
from functools import partial
from glob import glob
from git import Repo
from itertools import takewhile
import json
import os
from PIL import Image
from ply import lex, yacc
from pprint import pprint
from subprocess import Popen
import sys

REPO_DIR = 'overleaf_repos'

def overleaf2word(url,name=None,files=[]) :
    
    repo_root = url.split('/')[-1]

    if name is not None :
      # sanitize name
      replace = ' &;()/\\'
      repo_root = name.translate(str.maketrans(replace,'_'*len(replace))).lower()

    repo_dir = '{}/{}'.format(REPO_DIR,repo_root)
    
    if not os.path.exists(REPO_DIR) :
      os.mkdir(REPO_DIR)
    
    if not os.path.exists(repo_dir) :
        repo = Repo.clone_from(url,repo_dir)
    else :
        repo = Repo(repo_dir)
        repo.remotes['origin'].pull()
        
    # check for .bib file
    bibfn = glob(os.path.join(repo_dir,'*.bib'))
    if bibfn :
        bibfn = bibfn[0]
    else :
        bibfn = None
    for fn in files :
        tex_to_word(os.path.join(repo_dir,fn),repo_dir,bibfn)
        

##########################################################################################
# this is the ply tokenizer for latex
tokens = ('COMMAND','EQUATION','WORD','COMMENT','NEWLINE')

def t_COMMAND(t):
    r'\\(?!%)([a-zA-Z]+\*?)(?:\[([^]]+)\])?(?:{([^}]+)})?(\[[^]]+\])?'
    t.command, t.opts, t.args, t.post_opts = t.lexer.lexmatch.groups()[1:5]
    return t

t_COMMENT = r'(?<![\\])%.*'
def t_NEWLINE(t):
    r'\n'
    t.lexer.lineno += len(t.value)
    return t
t_WORD = r'[^\ \n]+'
def t_EQUATION(t) :
    r'[$][^$]+[$]'
    return t

t_ignore = ' '

def t_error(t) :
    print(t)
    return t

lexer = lex.lex()
##########################################################################################

Text = namedtuple('Text',['text','type','style','props'])
Word = partial(Text,type='word',style=None,props=None)

# for the life of me I can't figure out how to get python 2 and 3 support with
# this function
# you apparently can't call unicode(s,errors='replace') if s is a unicode object
# so you have to check for whether it is a unicode type object first, which
# doesn't work in python 3
# could someone else figure out how to fix this some day?
tou = str
if sys.version_info.major == 2:
    def wtf_unicode(text) :
        if isinstance(text,unicode) :
            return text
            
        # text can come in different encodings, coerce it to utf-8
        try :
            text = unicode(text,errors='replace')
        except TypeError as e :
            print('Error encoding unicode, pass:',e)
        return text
    tou = wtf_unicode
    
def add_run(par,words) :
    if words :
        # add a space at the end for funsies
        text = ' '.join(_.text for _ in words)
        text += ' '
        
        # convert to unicode
        text = tou(text)
        
        r = par.add_run(text,words[-1].style)
        if words[-1].props :
            for k,v in words[-1].props.items() :
                setattr(r,k,v)
            
def add_paragraph(doc,words) :
    '''Add words to doc as paragraph. Words is a list of Text namedtuples. The
    *type* field is used to collapse the words into runs with the same formatting.'''
    # collapse the words into lists of identical type to be turned into a
    # formatted run
    if words :
        par = doc.add_paragraph()
        curr_type = [words.pop(0)]
        while words:
            curr_text = words.pop(0)
            # curr_text is a different type than what we've seen already
            if curr_text.type != curr_type[-1].type : # new run
                add_run(par,curr_type)
                curr_type = [curr_text]
            else :
                curr_type.append(curr_text)
        add_run(par,curr_type)
    
def tex_to_word(tex_fn,repo_dir,bib_fn=None) :
    r"""Convert a LaTeX formatted file to docx format
    
    Parses ``tex_fn`` and converts text and some markup tags and environments
    into Word constructs. Creates a file with the same basename as ``tex_fn``
    but with ``.docx`` extension, e.g. ``main.tex -> main.docx``.
    
    If ``bib_fn`` is not provided, all ``\cite`` tags are replaced by parentheses,
    leaving keys as is. If ``bib_fn`` is provided, all ``\cite`` tags are replaced
    by <Author> <Year> formatted references, and if a ``\bibliography`` tag is
    present, a Reference section is formatted at the end of the document.
    
    :param tex_fn: path to LaTeX formatted file
    :param bib_fn: optional path to BibTeX formatted file containing citation
        information
    :return: nothing
    """
    
    print('\n-------------------------------------------------------------------')
    print(tex_fn)
        
    with open(tex_fn) as f :
        tex = f.read()
    parsed = lexer.input(tex)
    
    bibdb = None
    if bib_fn :
        with open(bib_fn) as f:
            bibtex_str = f.read()
        
        bibdb = bibtexparser.loads(bibtex_str)
        bibdb = {_['ID']:_ for _ in bibdb.entries}
    
    def is_heading(args) :
        return 'section' in args
        
    def get_heading_level(args) :
        return args.count('sub')+1
        
    heading_level = 1
    in_doc = False
    prev_token = None
    
    doc = docx.Document()
    text_started = False
    refs = set()
    words = []
    
    while True :
        tok = lexer.token()
        if not tok: break
        
        # handle commands, which control the structure of the document
        # and special elements like tables and figures (not yet implemented)
        if tok.type == 'COMMAND' :
            
            if tok.command == 'title' :
                doc.add_heading(tok.args,0)
            
            elif tok.command == 'begin' :
                
                # don't insert anything until we have seen a begin{document}
                if tok.args == 'document' :
                    in_doc = True
                    
                # other \begin's to be supported:
                # table
                # tabular
                # figure
                    
            # \section, \subsection, \subsubsection, etc
            elif is_heading(tok.command) :
                if words :
                    add_paragraph(doc,words)
                    words = []
                heading_level = get_heading_level(tok.command)
                doc.add_heading(tok.args,heading_level)
                
            # insert citation text
            elif tok.command == 'cite':
                ref_strs = tok.args
                if bibdb :
                    refids = tok.args.split(',')
                    refids = [_ for _ in refids if _]
                    refs.update(set(refids))
                    ref_strs = []
                    for refid in refids :
                        entry = bibdb[refid] 
                        author = entry.get('author',entry.get('title','')).split(',')[0]           
                        year = entry.get('year','')
                        ref_strs.append(' '.join([author,year]))
                    ref_strs = ','.join(ref_strs)
                citation = Text(
                    text=''.join(['(',ref_strs,')']),
                    type='cite',
                    style=None,
                    props=None
                )
                words.append(citation)
                
            elif tok.command == 'textbf' :
                bold = Text(
                    text=tok.args,
                    type='textbf',
                    style=None,
                    props={'bold':True}
                )
                words.append(bold)
                
            elif tok.command == 'textit' :
                italic = Text(
                    text=tok.args,
                    type='textit',
                    style=None,
                    props={'italic':True}
                )               
                words.append(italic)
                
            elif tok.command == 'clearpage' :
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                
            elif tok.command == 'includegraphics' :
                pic_path = os.path.join(repo_dir,tok.args)
                img = Image.open(pic_path)
                # calculate the image width in inches assuming 72 dpi
                # maximum 6 inches
                dpi = 72
                img_width = min(img.size[0]/72,6)
                
                doc.add_picture(pic_path,width=Inches(img_width))
                
            else :
                print('unrecognized command:',tok.command,tok.opts,tok.args)
        if tok.type == 'EQUATION' :
            print('found an equation',tok.value)
                    
        # regular text word
        if tok.type == 'WORD' :
            # replace escaped percents with literal percents
            tok.value = tok.value.replace(r'\%','%')

            text_started = True
            text = Word(text=tok.value)
            words.append(text)
            

        # if we hit two newlines in a row, create a new paragraph
        if tok.type == 'NEWLINE' and \
           prev_token and prev_token.type == 'NEWLINE' and \
           text_started :
               add_paragraph(doc,words)
               words = []
            
        prev_token = tok
            
    # do refs if there are refs
    
    if refs :
        doc.add_heading('References',heading_level)
        
        refs = sorted(list(refs))
        for i,refid in enumerate(refs) :
            ref = bibdb[refid]
            author = ''
            if 'author' in ref :
                author = ref['author'].split(' and ')
                author = author[0]+u' et al. '
            title = (tou(ref.get('title',''))
                .replace('{','')
                .replace('}','')
                .replace('\n',' ')
            )
            ref_words = [Word(text='{}. '.format(i+1)),
                Word(text=author),
                Word(text=title+u'. ')]
                
            def fmt(key,pref='',suff='') :
                if key in ref :
                    return Word(tou(pref+ref[key]+suff))
                
            ref_words.extend([
                fmt('journal',suff=u'. '),
                fmt('booktitle',suff=u'. '),
                fmt('volume',suff=u', '),
                fmt('pages',suff=u' '),
                fmt('year',pref=u'(',suff=u')'),
                fmt('howpublished',pref=u'(',suff=u')'),
                fmt('note'),
                Word(text=u'.')
            ])
            ref_words = [_ for _ in ref_words if _]
            add_paragraph(doc,ref_words)
                
    """
    [{'journal': 'Nice Journal',
      'comments': 'A comment',
      'pages': '12--23',
      'month': 'jan',
      'abstract': 'This is an abstract. This line should be long enough to test\nmultilines...',
      'title': 'An amazing title',
      'year': '2013',
      'volume': '12',
      'ID': 'Cesar2013',
      'author': 'Jean César',
      'keyword': 'keyword1, keyword2',
      'ENTRYTYPE': 'article'}]
    """

    # write out the doc
    basedir = os.path.dirname(tex_fn)
    basename, ext = os.path.splitext(os.path.basename(tex_fn))
    doc_fn = os.path.join(basedir,'{}.docx'.format(basename))
    doc.save(doc_fn)
