import bibtexparser
import docx
from ply import lex, yacc
from pprint import pprint

tokens = ('COMMAND','WORD','COMMENT','NEWLINE')

def t_COMMAND(t):
    r'\\([a-zA-Z]+)(?:\[([^]]+)\])?(?:{([^}]+)})?'
    t.command, t.opts, t.args = t.lexer.lexmatch.groups()[1:4]
    return t

t_COMMENT = r'%.*'
def t_NEWLINE(t):
    r'\n'
    t.lexer.lineno += len(t.value)
    return t
t_WORD = r'[^\ %\n]+'
t_ignore = ' '

def t_error(t) :
    print(t)
    return t

lexer = lex.lex()

test_dir = 'overleaf_repos/9997464gbzfdhbpqxsd/'
test_fn = '{}/main.tex'.format(test_dir)
with open(test_fn) as f :
    tex = f.read()
parsed = lexer.input(tex)

bib_fn = '{}/sample.bib'.format(test_dir)
with open(bib_fn) as f:
    bibtex_str = f.read()

bibdb = bibtexparser.loads(bibtex_str)
bibdb = {_['ID']:_ for _ in bibdb.entries}

def is_heading(args) :
    return 'section' in args
    
def get_heading_level(args) :
    return args.count('sub')+1
        
    
heading_level = 1
in_doc = False
prev_token = None

doc = docx.Document()
text_started = False
words = []

while True :
    tok = lexer.token()
    if not tok: break

    # handle commands, which control the structure of the document
    # and special elements like tables and figures (not yet implemented)
    if tok.type == 'COMMAND' :
        
        if tok.command == 'title' :
            doc.add_heading(tok.args,0)
        
        if tok.command == 'begin' :
            
            # don't insert anything until we have seen a begin{document}
            if tok.args == 'document' :
                in_doc = True
                
            # other \begin's to be supported:
            # table
            # tabular
            # figure
                
        # \section, \subsection, \subsubsection, etc
        if is_heading(tok.command) :
            if words :
                doc.add_paragraph(' '.join(words))
                words = []
            heading_level = get_heading_level(tok.command)
            doc.add_heading(tok.args,heading_level)
            
        # insert citation text
        if tok.command == 'cite':
            refids = tok.args.split(',')
            refids = [_ for _ in refids if _]
            ref_strs = []
            for refid in refids :
                entry = bibdb[refid] 
                author = entry.get('author',entry.get('title','')).split(',')[0]           
                year = entry.get('year','')
                ref_strs.append(' '.join([author,year]))
            ref_strs = ','.join(ref_strs)
            words.append(''.join(['(',ref_strs,')']))
                
    # regular text word
    if tok.type == 'WORD' :
        text_started = True
        words.append(tok.value)
        
    # if we hit two newlines in a row, create a new paragraph
    if tok.type == 'NEWLINE' and \
       prev_token and prev_token.type == 'NEWLINE' and \
       text_started :
            doc.add_paragraph(' '.join(words))
            words = []
        
    prev_token = tok
        
doc.save('{}/main.docx'.format(test_dir))
    